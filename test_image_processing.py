#!/usr/bin/env python3
"""
Test script to verify end-to-end IMAGE tag functionality.
"""

import os
import sys
import tempfile

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from report_compiler.document.placeholder_parser import PlaceholderParser
from report_compiler.document.docx_processor import DocxProcessor

def create_test_image():
    """Create a simple test image."""
    from PIL import Image
    
    # Create a simple 200x150 blue rectangle with text
    img = Image.new('RGB', (200, 150), color='blue')
    test_image_path = '/tmp/test_image_blue.png'
    img.save(test_image_path)
    return test_image_path

def create_test_docx_with_images():
    """Create a test DOCX with multiple IMAGE placeholders."""
    from docx import Document
    from docx.shared import Inches
    
    doc = Document()
    doc.add_heading('Test Document with IMAGE Tags', 0)
    
    doc.add_paragraph('This document contains different types of IMAGE placeholders:')
    
    # Test 1: Simple image placeholder
    doc.add_heading('Test 1: Simple Image', level=1)
    table1 = doc.add_table(rows=1, cols=1)
    table1.style = 'Table Grid'
    cell1 = table1.cell(0, 0)
    cell1.text = "[[IMAGE: /tmp/test_image_blue.png]]"
    table1.columns[0].width = Inches(3)
    table1.rows[0].height = Inches(2)
    
    # Test 2: Image with width parameter
    doc.add_heading('Test 2: Image with Width', level=1)
    table2 = doc.add_table(rows=1, cols=1)
    table2.style = 'Table Grid'
    cell2 = table2.cell(0, 0)
    cell2.text = "[[IMAGE: /tmp/test_image_blue.png, width=1.5in]]"
    table2.columns[0].width = Inches(4)
    table2.rows[0].height = Inches(1.5)
    
    # Test 3: Image with both width and height
    doc.add_heading('Test 3: Image with Width and Height', level=1)
    table3 = doc.add_table(rows=1, cols=1)
    table3.style = 'Table Grid'
    cell3 = table3.cell(0, 0)
    cell3.text = "[[IMAGE: /tmp/test_image_blue.png, width=2in, height=1in]]"
    table3.columns[0].width = Inches(3)
    table3.rows[0].height = Inches(1.5)
    
    test_docx_path = '/tmp/test_document_with_images.docx'
    doc.save(test_docx_path)
    return test_docx_path

def test_image_processing():
    """Test the complete image processing workflow."""
    print("Testing complete IMAGE processing workflow...\n")
    
    # Create test files
    test_image_path = create_test_image()
    test_docx_path = create_test_docx_with_images()
    output_docx_path = '/tmp/output_with_images.docx'
    
    print(f"Created test image: {test_image_path}")
    print(f"Created test document: {test_docx_path}")
    
    # Step 1: Parse placeholders
    print("\n1. Parsing placeholders...")
    parser = PlaceholderParser()
    placeholders = parser.find_all_placeholders(test_docx_path)
    
    print(f"   Found {placeholders['total']} total placeholders")
    for i, placeholder in enumerate(placeholders['table']):
        print(f"   Table {i}: {placeholder['subtype']} - {placeholder['file_path']}")
    
    # Step 2: Process the document
    print("\n2. Processing document...")
    processor = DocxProcessor()
    metadata = processor.create_modified_docx(test_docx_path, placeholders, output_docx_path)
    
    if metadata is not None:
        print(f"   ✓ Document processed successfully")
        print(f"   ✓ Output saved to: {output_docx_path}")
        print(f"   ✓ Table metadata: {metadata}")
        
        # Check if output file exists and has reasonable size
        if os.path.exists(output_docx_path):
            size = os.path.getsize(output_docx_path)
            print(f"   ✓ Output file size: {size} bytes")
            
            # Try to open and verify the output document
            try:
                from docx import Document
                output_doc = Document(output_docx_path)
                print(f"   ✓ Output document can be opened")
                print(f"   ✓ Number of tables in output: {len(output_doc.tables)}")
                
                # Check if images were actually inserted
                for i, table in enumerate(output_doc.tables):
                    if len(table.rows) == 1 and len(table.columns) == 1:
                        cell = table.cell(0, 0)
                        # Check if the cell has any inline shapes (images)
                        inline_shapes = cell.paragraphs[0]._element.xpath('.//a:blip')
                        if inline_shapes:
                            print(f"   ✓ Table {i} contains an image")
                        else:
                            print(f"   ? Table {i} content: '{cell.text[:50]}'")
                            
            except Exception as e:
                print(f"   ✗ Error verifying output document: {e}")
        else:
            print(f"   ✗ Output file was not created")
    else:
        print(f"   ✗ Document processing failed")

if __name__ == "__main__":
    test_image_processing()