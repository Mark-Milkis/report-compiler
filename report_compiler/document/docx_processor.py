"""
DOCX document processing and modification utilities.
"""

import os
import shutil
from typing import Dict, List, Any, Optional
from docx import Document
from ..core.config import Config
from ..utils.logging_config import get_docx_logger


class DocxProcessor:
    """Handles DOCX document modification and marker insertion."""
    
    def __init__(self, input_path: str):
        self.input_path = input_path
        self.doc = None
        self.placeholders = {}
        self.table_metadata = {}
        self.logger = get_docx_logger()
    
    def load_document(self) -> None:
        """Load the DOCX document."""
        self.doc = Document(self.input_path)
    
    def create_modified_document(self, placeholders: Dict[str, List[Dict]], 
                                output_path: str) -> bool:
        """
        Create a modified DOCX document with markers inserted.
        
        Args:
            placeholders: Dictionary containing table and paragraph placeholders
            output_path: Path for the modified document
            
        Returns:
            bool: True if successful, False otherwise
        """
        try:
            if self.doc is None:
                self.load_document()
            
            self.placeholders = placeholders
            self.logger.info("🔧 PHASE 2: Modifying document...")
            self.logger.info("🔧 Creating modified document...")
            
            # Process merge placeholders first (paragraph-based)
            if placeholders['paragraph']:
                self._process_merge_placeholders()
            
            # Process overlay placeholders (table-based)
            if placeholders['table']:
                self._process_overlay_placeholders()
            
            # Save the modified document
            self.doc.save(output_path)
            self.logger.info("✅ Document modification complete")
            
            return True
            
        except Exception as e:
            self.logger.error("❌ Error creating modified document: %s", e, exc_info=True)
            return False
    
    def _process_merge_placeholders(self) -> None:
        """Process paragraph-based merge placeholders."""
        merge_placeholders = self.placeholders['paragraph']
        self.logger.info("📄 Processing %d merge placeholders...", len(merge_placeholders))
        
        for idx, placeholder in enumerate(merge_placeholders, 1):
            page_count = placeholder.get('page_count', 0)
            para_idx = placeholder['paragraph_index']
            self.logger.info("   📄 Processing merge placeholder #%d:", idx)
            self.logger.info("      • Paragraph %d, %d pages", para_idx, page_count)
            
            # Generate marker
            marker = Config.get_merge_marker(idx)
            self.logger.info("      • Marker: %s", marker)
            
            # Find the paragraph and replace its content
            if para_idx < len(self.doc.paragraphs):
                paragraph = self.doc.paragraphs[para_idx]
                
                # Replace placeholder text with marker
                paragraph.clear()
                paragraph.add_run(marker)
                  # Add page break after marker
                from docx.enum.text import WD_BREAK
                paragraph.add_run().add_break(WD_BREAK.PAGE)
                
                self.logger.info("      ✅ Added marker and page break (no placeholder pages)")
            else:
                self.logger.warning("      ⚠️ Paragraph index %d out of range", para_idx)
    
    def _process_overlay_placeholders(self) -> None:
        """Process table-based overlay placeholders."""
        overlay_placeholders = self.placeholders['table']
        self.logger.info("📦 Processing %d overlay placeholders...", len(overlay_placeholders))
        
        self.logger.info("   📋 Table-based overlay placeholder processing:")
        
        for idx, placeholder in enumerate(overlay_placeholders):
            table_idx = placeholder['table_index']
            page_count = placeholder.get('page_count', 1)
            
            self.logger.info("      • Processing table placeholder #%d:", idx + 1)
            self.logger.info("         • Table %d, %d pages", table_idx, page_count)
            
            # Get stored table dimensions
            dimensions = self._extract_table_dimensions(placeholder)
            self.logger.info("         • Using stored table dimensions: %.2f x %.2f inches", 
                           dimensions['width_inches'], dimensions['height_inches'])
            
            # Convert to points for PDF processing
            table_width_pts = dimensions['width_inches'] * 72
            table_height_pts = dimensions['height_inches'] * 72
            self.logger.info("         • Final table dimensions: %.1f x %.1f points = %.2f x %.2f inches",
                           table_width_pts, table_height_pts, dimensions['width_inches'], dimensions['height_inches'])
            
            # Store metadata for later use
            self.table_metadata[table_idx] = {
                'width_pts': table_width_pts,
                'height_pts': table_height_pts
            }
            
            # Process the table
            if table_idx < len(self.doc.tables):
                table = self.doc.tables[table_idx]
                
                if page_count > 1:
                    self.logger.info("         📋 Multi-page PDF detected (%d pages), replicating cells...", page_count)
                    self._replicate_table_cells(table, table_idx, page_count)
                else:
                    # Single page - just replace with marker
                    cell = table.rows[0].cells[0]
                    marker_text = Config.get_overlay_marker(table_idx)
                    
                    cell.text = "" # Clear all content, including paragraphs
                    # Get the first paragraph or add one if none exist
                    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                    # The legacy code also sets alignment here. If needed, add:
                    # from docx.enum.text import WD_ALIGN_PARAGRAPH
                    # paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.add_run(marker_text)
                
                self.logger.info("         ✅ Table %d updated with overlay marker and dimensions", table_idx)
            else:
                self.logger.warning("         ⚠️ Table index %d out of range", table_idx)
    
    def _replicate_table_cells(self, table, table_idx: int, page_count: int) -> None:
        """Replicate table cells for multi-page overlays."""
        # First cell gets the main marker
        first_cell = table.rows[0].cells[0]
        main_marker_text = Config.get_overlay_marker(table_idx)
        
        first_cell.text = "" # Clear all content
        # Get the first paragraph or add one if none exist
        paragraph = first_cell.paragraphs[0] if first_cell.paragraphs else first_cell.add_paragraph()
        # The legacy code also sets alignment here. If needed, add:
        # from docx.enum.text import WD_ALIGN_PARAGRAPH
        # paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.add_run(main_marker_text)
        
        # Add additional rows for each additional page
        created_cells = 0
        for page_num in range(2, page_count + 1):
            # Add a new row
            new_row = table.add_row()
            
            # Attempt to set row height to match the original first row's height
            try:
                if table.rows[0].height is not None:
                    new_row.height = table.rows[0].height
            except Exception as e:
                self.logger.warning("           ⚠️ Could not set row height for replicated cell (table_idx: %d, row: %d): %s", 
                                  table_idx, page_num, e)
                pass # Continue without setting height, similar to legacy try-except pass

            new_cell = new_row.cells[0]
            
            # Set marker for this page
            page_marker_text = Config.get_overlay_marker(table_idx, page_num)
            new_cell.text = "" # Clear all content
            # Get the first paragraph or add one if none exist
            paragraph_repl = new_cell.paragraphs[0] if new_cell.paragraphs else new_cell.add_paragraph()
            # The legacy code also sets alignment here. If needed, add:
            # from docx.enum.text import WD_ALIGN_PARAGRAPH
            # paragraph_repl.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_repl.add_run(page_marker_text)
            
            self.logger.info("           ✅ Created table row %d with marker: %s", page_num, page_marker_text)
            created_cells += 1
        
        self.logger.info("         ✅ Created %d additional cells", created_cells)
    
    def _extract_table_dimensions(self, placeholder: Dict[str, Any]) -> Dict[str, float]:
        """Extract table dimensions from placeholder metadata."""
        dimensions = {}
        
        # Try to get width
        if 'width_inches' in placeholder:
            dimensions['width_inches'] = placeholder['width_inches']
        elif 'column_width_inches' in placeholder:
            dimensions['width_inches'] = placeholder['column_width_inches']
        else:
            # Default width if not found
            dimensions['width_inches'] = 7.5
            self.logger.warning("         ⚠️ No width found, using default: %s inches", dimensions['width_inches'])
        
        # Try to get height
        if 'height_inches' in placeholder:
            dimensions['height_inches'] = placeholder['height_inches']
        elif 'row_height_inches' in placeholder:
            dimensions['height_inches'] = placeholder['row_height_inches']
        else:
            # Default height if not found
            dimensions['height_inches'] = 4.0
            self.logger.warning("         ⚠️ No height found, using default: %s inches", dimensions['height_inches'])
        
        return dimensions
    
    def get_table_metadata(self) -> Dict[int, Dict[str, float]]:
        """Get stored table metadata for PDF processing."""
        return self.table_metadata
