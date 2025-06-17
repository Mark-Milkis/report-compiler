"""
Placeholder detection and parsing for DOCX documents.
"""

import re
from typing import Dict, List, Any, Optional
from docx import Document
from ..core.config import Config
from ..utils.logging_config import get_module_logger


class PlaceholderParser:
    """Handles detection and parsing of PDF placeholders in DOCX documents."""
    
    def __init__(self):
        self.overlay_regex = Config.OVERLAY_REGEX
        self.insert_regex = Config.INSERT_REGEX
        self.logger = get_module_logger(__name__)
        
        # Cache for document parsing
        self._doc = None
        self._doc_path = None
    
    def find_all_placeholders(self, docx_path: str) -> Dict[str, List[Dict]]:
        """
        Find all placeholders in a DOCX document.
        
        Args:
            docx_path: Path to DOCX document
            
        Returns:
            Dict with 'table' and 'paragraph' placeholder lists
        """
        self._load_document(docx_path)
        
        table_placeholders = self._find_table_placeholders()
        paragraph_placeholders = self._find_paragraph_placeholders()
        
        return {
            'table': table_placeholders,
            'paragraph': paragraph_placeholders,
            'total': len(table_placeholders) + len(paragraph_placeholders)
        }
    
    def _load_document(self, docx_path: str) -> None:
        """Load document if not already loaded or path changed."""
        if self._doc is None or self._doc_path != docx_path:
            self._doc = Document(docx_path)
            self._doc_path = docx_path
    
    def _find_table_placeholders(self) -> List[Dict]:
        """
        Find PDF placeholders inside single-cell tables (overlay type).
        
        Returns:
            List of dictionaries containing table placeholder info
        """
        placeholders = []
        
        try:
            for table_idx, table in enumerate(self._doc.tables):
                rows = len(table.rows)
                cols = len(table.columns)
                
                # Only consider single-cell tables for overlay inserts
                if rows == 1 and cols == 1:
                    cell = table.rows[0].cells[0]
                    cell_text = cell.text.strip()
                    
                    # Check if this cell contains an OVERLAY placeholder
                    match = self.overlay_regex.search(cell_text)
                    if match:
                        pdf_path_raw = match.group(1).strip()
                        params_string = match.group(2)                        
                        # Parse parameters
                        params = self._parse_overlay_parameters(params_string)
                        
                        self.logger.info("   📋 Found table OVERLAY placeholder #%d:", len(placeholders)+1)
                        self.logger.info("      • Raw path: %s", pdf_path_raw)
                        if params['page']:
                            self.logger.info("      • Page specification: page=%s", params['page'])
                        self.logger.info("      • Content cropping: %s", 'enabled' if params['crop'] else 'disabled')
                        self.logger.info("      • Table index: %d", table_idx)
                        self.logger.info("      • Table type: Single-cell (1x1)")
                        self.logger.info("      • Cell text: '%s'", cell_text)
                        
                        # Get table dimensions
                        dimensions = self._get_table_dimensions(table, table_idx)
                        
                        table_info = {
                            'type': 'overlay',
                            'pdf_path_raw': pdf_path_raw,
                            'page_spec': params['page'],
                            'crop_enabled': params['crop'],
                            'table_index': table_idx,
                            'table_text': cell_text,
                            'source': f'table_{table_idx}',
                            'insert_method': 'table'
                        }
                        
                        if dimensions:
                            table_info.update(dimensions)
                            if 'width_inches' in dimensions and 'height_inches' in dimensions:
                                self.logger.info("      • Dimensions: %.2f\" x %.2f\"", dimensions['width_inches'], dimensions['height_inches'])
                            else:
                                self.logger.info("      • Dimensions: %s", dimensions)
                        else:
                            self.logger.warning("      • ⚠️ Could not determine table dimensions")
                        
                        placeholders.append(table_info)
                
                else:                    # Multi-cell tables: scan but don't classify as overlay
                    has_insert = False
                    for row in table.rows:
                        for cell in row.cells:
                            if (self.overlay_regex.search(cell.text) or 
                                self.insert_regex.search(cell.text)):
                                has_insert = True
                                break
                        if has_insert:
                            break
                    
                    if has_insert:
                        self.logger.warning("   ⚠️  Multi-cell table #%d (%dx%d) contains placeholder but skipped (not overlay type)", table_idx, rows, cols)
        
        except Exception as e:
            self.logger.error("   ❌ Error scanning for table placeholders: %s", e, exc_info=True)
        
        self.logger.info("   ✅ Found %d table placeholders", len(placeholders))
        return placeholders
    
    def _find_paragraph_placeholders(self) -> List[Dict]:
        """
        Find PDF placeholders in regular paragraphs (merge type).
        
        Returns:
            List of dictionaries containing paragraph placeholder info
        """
        placeholders = []
        
        try:
            for para_idx, paragraph in enumerate(self._doc.paragraphs):
                para_text = paragraph.text.strip()
                
                # Look for INSERT placeholders (merge type)
                match = self.insert_regex.search(para_text)
                if match:
                    pdf_path_raw = match.group(1).strip()
                    page_spec = match.group(2)  # Optional page specification                    
                    self.logger.info("   📄 Found paragraph INSERT placeholder #%d:", len(placeholders)+1)
                    self.logger.info("      • Raw path: %s", pdf_path_raw)
                    if page_spec:
                        self.logger.info("      • Page specification: %s", page_spec)
                    self.logger.info("      • Paragraph index: %d", para_idx)
                    
                    placeholder_info = {
                        'type': 'merge',
                        'pdf_path_raw': pdf_path_raw,
                        'page_spec': page_spec,
                        'paragraph_index': para_idx,
                        'paragraph_text': para_text,
                        'source': f'paragraph_{para_idx}',
                        'insert_method': 'merge'
                    }
                    
                    placeholders.append(placeholder_info)
        
        except Exception as e:
            self.logger.error("   ❌ Error scanning for paragraph placeholders: %s", e, exc_info=True)
        
        self.logger.info("   ✅ Found %d paragraph placeholders", len(placeholders))
        return placeholders
    
    def _parse_overlay_parameters(self, params_string: Optional[str]) -> Dict[str, Any]:
        """
        Parse overlay parameters from the placeholder text.
        
        Args:
            params_string: Parameter string after the path (can be None)
            
        Returns:
            Dict with parsed parameters
        """
        result = {
            'page': None,
            'crop': Config.DEFAULT_CROP_ENABLED
        }
        
        if not params_string:
            return result
        
        # Split parameters by comma
        params = [p.strip() for p in params_string.split(',')]
        
        for param in params:
            if '=' in param:
                key, value = param.split('=', 1)
                key = key.strip().lower()
                value = value.strip()
                
                if key == 'page':
                    result['page'] = value
                elif key == 'crop':
                    result['crop'] = value.lower() in ('true', '1', 'yes', 'on', 'enabled')
            else:
                # Assume it's a page specification if no key
                if param and not result['page']:
                    result['page'] = param
        
        return result
    
    def _get_table_dimensions(self, table, table_idx: int) -> Optional[Dict[str, Any]]:
        """
        Extract dimensions and position information from a table element.
        
        Args:
            table: python-docx table object
            table_idx: Index of the table
            
        Returns:
            Dict with width/height information or None
        """
        try:
            dimensions = {}
            
            if hasattr(table, '_tbl'):
                tbl_element = table._tbl
                tbl_pr = tbl_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
                
                if tbl_pr is not None:
                    # Look for width information
                    tbl_w = tbl_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
                    if tbl_w is not None:
                        width_type = tbl_w.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'unknown')
                        width_val = tbl_w.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', 'unknown')
                        
                        if width_type == 'dxa' and width_val != 'unknown':
                            # Convert from twentieths of a point to inches
                            width_inches = int(width_val) / 1440.0
                            dimensions['width_inches'] = width_inches
                            dimensions['width_type'] = width_type
                            dimensions['width_raw'] = width_val
            
            # Try to get column width (alternative method)
            if 'width_inches' not in dimensions:
                try:
                    first_cell = table.rows[0].cells[0]
                    if hasattr(first_cell, 'width') and first_cell.width:
                        # Width is in EMU (English Metric Units), convert to inches
                        width_inches = first_cell.width.inches
                        dimensions['column_width_inches'] = width_inches
                except:
                    pass
            
            # Try to get row height
            try:
                first_row = table.rows[0]
                if hasattr(first_row, 'height') and first_row.height:
                    height_inches = first_row.height.inches
                    dimensions['row_height_inches'] = height_inches
            except:
                pass
              # Store coordinate metadata for later use
            self.logger.debug("      📍 Stored coordinate metadata for table %d", table_idx)
            if dimensions:
                self.logger.debug("      • Dimensions: %s", dimensions)
            
            return dimensions if dimensions else None
            
        except Exception as e:
            self.logger.warning("      ⚠️ Error extracting table dimensions: %s", e)
            return None
