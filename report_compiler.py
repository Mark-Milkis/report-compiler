#!/usr/bin/env python3
"""
Python PDF Report Compiler using DOCX Modification and PDF Overlay

This script automates the creation of a final PDF report by combining a main Word document 
with multiple PDF appendices. It supports two types of placeholders:
- Table-based overlays that overlay content onto existing pages
- Paragraph-based merges that insert PDF pages after markers with page breaks
"""

import os
import re
import argparse
import win32com.client as win32
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_BREAK
import fitz  # PyMuPDF
import time


class ReportCompiler:
    """
    A class to compile Word documents with PDF appendices into a single PDF report.
      The process involves:
    1. Finding PDF insertion placeholders in the Word document
    2. Modifying the document to create markers and page breaks as needed
    3. Converting the modified document to PDF
    4. Processing the PDF insertions (overlays and merges)
    
    Supported placeholder types:
    - Table-based overlays: Single-cell tables containing [[OVERLAY: path.pdf, page=5]]
    - Paragraph-based merges: Regular paragraphs containing [[INSERT: path.pdf]]
    """
    
    def __init__(self, input_docx_path, final_pdf_path, keep_temp=False):
        """
        Initialize the ReportCompiler with input and output paths.
        
        Args:
            input_docx_path (str): Absolute path to the source .docx file
            final_pdf_path (str): Absolute path for the final output .pdf file
            keep_temp (bool): Whether to keep temporary files for debugging
        """
        self.input_docx_path = os.path.abspath(input_docx_path)
        self.final_pdf_path = os.path.abspath(final_pdf_path)
        self.keep_temp = keep_temp
          # Store table coordinate metadata for overlay positioning
        self.table_coordinates = {}  # Maps table index to position/size info
        
        # Create unique temporary file names using timestamp
        timestamp = str(int(time.time() * 1000))
        base_dir = os.path.dirname(self.input_docx_path)
        self.temp_docx_path = os.path.join(base_dir, f"~temp_modified_report_{timestamp}.docx")
        self.temp_pdf_path = os.path.join(base_dir, f"~temp_base_{timestamp}.pdf")        # Compile regex patterns for finding placeholders
        # OVERLAY format: [[OVERLAY: path.pdf, page=5]] or [[OVERLAY: path.pdf]]
        # INSERT format: [[INSERT: path.pdf]] or [[INSERT: path.pdf:1-3]]
        # Note: Handle Windows paths with drive letters like C:\path
        self.overlay_regex = re.compile(r"\[\[OVERLAY:\s*([^,\]]+?)(?:,\s*page=([^\]]+))?\s*\]\]")
        self.insert_regex = re.compile(r"\[\[INSERT:\s*(.+?)(?::([^:\\\/\]]+))?\s*\]\]")
        
        print(f"Input DOCX: {self.input_docx_path}")
        print(f"Output PDF: {self.final_pdf_path}")
        print(f"Temp DOCX: {self.temp_docx_path}")
        print(f"Temp PDF: {self.temp_pdf_path}")
    
    def _parse_page_specification(self, page_spec):
        """
        Parse page specifications from INSERT statements.
        
        Supported formats:
        - None or empty: All pages
        - "5": Single page (5)
        - "1-3": Range of pages (1, 2, 3)
        - "2-": Pages from 2 to end
        - "1,3,5": Specific pages (1, 3, 5)
        - "1-3,7,9-": Mixed specification
        
        Args:
            page_spec (str): Page specification string
            
        Returns:
            None: Process all pages
            dict: Dictionary with 'pages' list and optional 'open_range_start'
        """
        if not page_spec or not page_spec.strip():
            return None  # Process all pages
        
        page_spec = page_spec.strip()
        selected_pages = set()
        
        try:
            # Split by commas for multiple specifications
            parts = [part.strip() for part in page_spec.split(',')]
            
            for part in parts:
                if '-' in part:
                    # Range specification
                    if part.endswith('-'):
                        # Open-ended range like "2-"
                        start = int(part[:-1]) - 1  # Convert to 0-based
                        selected_pages.add(('range_from', start))
                    else:
                        # Closed range like "1-3"
                        start_str, end_str = part.split('-', 1)
                        start = int(start_str) - 1  # Convert to 0-based
                        end = int(end_str) - 1      # Convert to 0-based
                        selected_pages.update(range(start, end + 1))
                else:
                    # Single page
                    page = int(part) - 1  # Convert to 0-based
                    selected_pages.add(page)
            
            # Handle open-ended ranges
            final_pages = []
            open_range_start = None
            
            for item in selected_pages:
                if isinstance(item, tuple) and item[0] == 'range_from':
                    open_range_start = item[1]
                else:
                    final_pages.append(item)
            
            # Sort the pages
            final_pages = sorted(final_pages)
            
            print(f"        📄 Page specification '{page_spec}' parsed to: {[p+1 for p in final_pages]}" + 
                  (f" plus {open_range_start+1}+" if open_range_start is not None else ""))
            
            return {
                'pages': final_pages,
                'open_range_start': open_range_start
            }
            
        except (ValueError, IndexError) as e:
            print(f"        ❌ Invalid page specification '{page_spec}': {e}")
            print(f"        📝 Using all pages instead")
            return None

    def run(self):
        """
        Main public method that executes the entire workflow.
        
        The workflow:
        1. Find placeholders and modify the DOCX
        2. Convert modified DOCX to PDF
        3. Overlay appendix PDFs onto the base PDF
        4. Clean up temporary files
        """
        try:
            print("\n=== Starting Report Compilation ===")
            
            # Step 1: Find placeholders and modify DOCX
            print("\nStep 1: Analyzing document for PDF placeholders...")
            modified_doc, placeholders = self._find_placeholders_and_modify_docx()
            
            if not placeholders:
                print("No PDF placeholders found. Converting original document to PDF...")
                self._convert_docx_to_pdf(self.input_docx_path, self.final_pdf_path)
                print(f"✓ Conversion complete: {self.final_pdf_path}")
                return
            
            print(f"Found {len(placeholders)} PDF placeholder(s)")
              # Step 2: Save modified document and convert to PDF
            print("\nStep 2: Saving modified document...")
            modified_doc.save(self.temp_docx_path)
            print(f"✓ Modified document saved: {self.temp_docx_path}")
            
            print("\nStep 3: Converting modified document to PDF...")
            self._convert_docx_to_pdf(self.temp_docx_path, self.temp_pdf_path)
            print(f"✓ Base PDF created: {self.temp_pdf_path}")
              # Step 4: Process PDF insertions by type
            print("\nStep 4: Processing PDF insertions...")
            overlay_placeholders = [p for p in placeholders if p.get('type') == 'overlay']
            merge_placeholders = [p for p in placeholders if p.get('type') == 'merge']
            
            # Process merges first (they modify page structure)
            if merge_placeholders:
                print(f"   • Processing {len(merge_placeholders)} paragraph-based merge(s)...")
                self._insert_pdfs(merge_placeholders)
                print(f"   ✓ Merge processing complete")
            
            # Process overlays second (they overlay on existing pages)
            if overlay_placeholders:
                print(f"   • Processing {len(overlay_placeholders)} table-based overlay(s)...")
                # If merges were processed, work from the merged PDF, otherwise from base PDF
                source_pdf = self.final_pdf_path if merge_placeholders else self.temp_pdf_path
                self._overlay_pdfs(overlay_placeholders, source_pdf)
                print(f"   ✓ Overlay processing complete")
            
            # If no processing occurred, copy base PDF to final
            if not overlay_placeholders and not merge_placeholders:
                print("   • No placeholders to process")
                import shutil
                shutil.copy2(self.temp_pdf_path, self.final_pdf_path)
                print(f"   ✓ Final PDF copied: {self.final_pdf_path}")
            
            print(f"\n✓ Final PDF created: {self.final_pdf_path}")
            print("\n=== Report Compilation Complete ===")
            
        finally:
            # Always clean up temporary files
            self._cleanup()
    
    def _find_placeholders_and_modify_docx(self):
        """
        Find PDF placeholders and sort them into two types: overlay and merge insertions.
        
        Overlay insertions: Placeholders found inside single-cell tables - use table size/position
        Merge insertions: Placeholders found in regular paragraphs - full page inserts
        
        Returns:
            tuple: (modified_doc, placeholders_list)
                - modified_doc: The modified docx.Document object
                - placeholders_list: List of dictionaries containing placeholder info
        """
        print("🔍 PHASE 1: Scanning document for INSERT placeholders...")
        
        # Step 1: Scan for table placeholders (overlay type)
        print("\n📋 Scanning tables for overlay placeholders...")
        table_placeholders = self._find_table_placeholders()
        
        # Step 2: Scan for paragraph placeholders (merge type)
        print("\n📄 Scanning paragraphs for merge placeholders...")
        paragraph_placeholders = self._find_paragraph_placeholders()
        
        # Step 3: Validate and count pages for all placeholders
        print(f"\n✅ VALIDATION: Found {len(table_placeholders)} table + {len(paragraph_placeholders)} paragraph placeholders")
        all_placeholders = table_placeholders + paragraph_placeholders
        validated_placeholders = self._validate_pdf_references(all_placeholders)
        
        # Step 4: Sort placeholders by type for processing
        overlay_placeholders = [p for p in validated_placeholders if p['type'] == 'overlay']
        merge_placeholders = [p for p in validated_placeholders if p['type'] == 'merge']
        
        print(f"\n📊 PROCESSING SUMMARY:")
        print(f"   • Overlay insertions (table-based): {len(overlay_placeholders)}")
        print(f"   • Merge insertions (paragraph-based): {len(merge_placeholders)}")
        
        if not validated_placeholders:
            print("⚠️  No valid placeholders found - document will be converted as-is")
            return Document(self.input_docx_path), []
        
        # Step 5: Modify document based on placeholder types
        print(f"\n🔧 PHASE 2: Modifying document...")
        modified_doc = self._create_modified_document(overlay_placeholders, merge_placeholders)
        
        return modified_doc, validated_placeholders
    
    def _find_table_placeholders(self):
        """
        Find PDF placeholders inside single-cell tables (overlay type).
        
        Single-cell tables containing INSERT placeholders are considered overlay inserts
        because they can be precisely positioned and sized.
        
        Returns:
            list: List of dictionaries containing table placeholder info
        """
        placeholders = []
        
        try:
            doc = Document(self.input_docx_path)
            
            for table_idx, table in enumerate(doc.tables):
                rows = len(table.rows)
                cols = len(table.columns)
                
                # Only consider single-cell tables for overlay inserts
                if rows == 1 and cols == 1:
                    cell = table.rows[0].cells[0]
                    cell_text = cell.text.strip()                    # Check if this cell contains an OVERLAY placeholder
                    match = self.overlay_regex.search(cell_text)
                    if match:
                        pdf_path_raw = match.group(1).strip()
                        page_spec = match.group(2)  # Page specification (could be None)
                        
                        print(f"   📋 Found table OVERLAY placeholder #{len(placeholders)+1}:")
                        print(f"      • Raw path: {pdf_path_raw}")
                        if page_spec:
                            print(f"      • Page specification: page={page_spec}")
                        print(f"      • Table index: {table_idx}")
                        print(f"      • Table type: Single-cell (1x1)")
                        print(f"      • Cell text: '{cell_text}'")
                          # Try to get table dimensions
                        dimensions = self._get_table_dimensions(table, table_idx)
                        
                        table_info = {
                            'type': 'overlay',
                            'pdf_path_raw': pdf_path_raw,
                            'page_spec': page_spec,
                            'table_index': table_idx,
                            'table_text': cell_text,
                            'source': f'table_{table_idx}',
                            'insert_method': 'table'
                        }
                        
                        if dimensions:
                            table_info.update(dimensions)
                            if 'width_inches' in dimensions and 'height_inches' in dimensions:
                                print(f"      • Dimensions: {dimensions['width_inches']:.2f}\" x {dimensions['height_inches']:.2f}\"")
                            else:
                                print(f"      • Dimensions: {dimensions}")
                        else:
                            print(f"      • ⚠️ Could not determine table dimensions")
                        
                        placeholders.append(table_info)
                
                else:                    # Multi-cell tables: scan all cells but don't classify as overlay
                    # This prevents inline text within tables from being misclassified
                    has_insert = False
                    for row in table.rows:
                        for cell in row.cells:
                            if (self.overlay_regex.search(cell.text) or 
                                self.insert_regex.search(cell.text)):
                                has_insert = True
                                break
                        if has_insert:
                            break
                    
                    if has_insert:
                        print(f"   ⚠️  Multi-cell table #{table_idx} ({rows}x{cols}) contains placeholder but skipped (not overlay type)")
        
        except Exception as e:
            print(f"   ❌ Error scanning for table placeholders: {e}")
        
        print(f"   ✅ Found {len(placeholders)} table placeholders")
        return placeholders
    
    def _get_table_dimensions(self, table, table_idx):
        """
        Extract dimensions and position information from a table element.
        Enhanced for Word-to-PDF coordinate mapping.
        
        Returns:
            dict: Dictionary with width/height information or None
        """
        try:
            # Method 1: Check table preferred width
            dimensions = {}
            
            if hasattr(table, '_tbl'):
                tbl_element = table._tbl
                tbl_pr = tbl_element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
                if tbl_pr is not None:
                    # Look for width information
                    tbl_w = tbl_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
                    if tbl_w is not None:
                        width_type = tbl_w.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'unknown')
                        width_val = tbl_w.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w', 'unknown')
                        
                        if width_type == 'dxa' and width_val != 'unknown':
                            # Convert from twentieths of a point to inches
                            width_inches = int(width_val) / 1440.0
                            dimensions['width_inches'] = width_inches
                            dimensions['width_type'] = width_type
                            dimensions['width_raw'] = width_val
                    
                    # Look for table positioning information
                    tbl_pos = tbl_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblpPr')
                    if tbl_pos is not None:
                        # Table has absolute positioning
                        h_anchor = tbl_pos.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}horzAnchor', 'margin')
                        v_anchor = tbl_pos.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vertAnchor', 'margin')
                        tbl_px = tbl_pos.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblpX')
                        tbl_py = tbl_pos.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblpY')
                        
                        dimensions['has_absolute_position'] = True
                        dimensions['h_anchor'] = h_anchor
                        dimensions['v_anchor'] = v_anchor
                        if tbl_px:
                            dimensions['pos_x_twips'] = int(tbl_px)
                            dimensions['pos_x_inches'] = int(tbl_px) / 1440.0
                        if tbl_py:
                            dimensions['pos_y_twips'] = int(tbl_py)
                            dimensions['pos_y_inches'] = int(tbl_py) / 1440.0
                    
                    # Look for table margins/indentation
                    tbl_ind = tbl_pr.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblInd')
                    if tbl_ind is not None:
                        ind_type = tbl_ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type', 'dxa')
                        ind_val = tbl_ind.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}w')
                        if ind_val and ind_type == 'dxa':
                            dimensions['indent_twips'] = int(ind_val)
                            dimensions['indent_inches'] = int(ind_val) / 1440.0
            
            # Method 2: Check column widths
            if table.columns:
                try:
                    first_col = table.columns[0]
                    if hasattr(first_col, 'width') and first_col.width:
                        col_width_inches = first_col.width.inches
                        dimensions['column_width_inches'] = col_width_inches
                except:
                    pass
            
            # Method 3: Try to get row height
            if table.rows:
                try:
                    first_row = table.rows[0]
                    if hasattr(first_row, 'height') and first_row.height:
                        row_height_inches = first_row.height.inches
                        dimensions['row_height_inches'] = row_height_inches
                except:
                    pass
              # Store coordinate metadata for precise overlay positioning
            if dimensions:
                self.table_coordinates[table_idx] = {
                    'dimensions': dimensions,
                    'word_table_index': table_idx,
                    'extraction_method': 'word_xml_analysis'
                }
                print(f"      📍 Stored coordinate metadata for table {table_idx}")
            
            return dimensions if dimensions else None
            
        except Exception as e:
            print(f"      ⚠️ Error getting table dimensions: {e}")
            return None
    
    def _find_paragraph_placeholders(self):
        """
        Find PDF placeholders in regular document paragraphs (merge type).
          Returns:
            list: List of dictionaries containing paragraph placeholder info
        """
        placeholders = []
        doc = Document(self.input_docx_path)
        for para_idx, paragraph in enumerate(doc.paragraphs):
            match = self.insert_regex.search(paragraph.text)
            if match:
                pdf_path_raw = match.group(1).strip()
                page_spec = match.group(2)  # Page specification (could be None)
                
                print(f"   📄 Found paragraph INSERT placeholder #{len(placeholders)+1}:")
                print(f"      • Raw path: {pdf_path_raw}")
                if page_spec:
                    print(f"      • Page specification: {page_spec}")
                print(f"      • Paragraph index: {para_idx}")
                
                placeholders.append({
                    'type': 'merge',
                    'pdf_path_raw': pdf_path_raw,
                    'page_spec': page_spec,
                    'paragraph_index': para_idx,
                    'source': f'paragraph_{para_idx}'
                })
        
        print(f"   ✅ Found {len(placeholders)} paragraph placeholders")
        return placeholders
    
    def _validate_pdf_references(self, placeholders):
        """
        Validate PDF file existence and count pages for all placeholders.
        
        Args:
            placeholders (list): List of placeholder dictionaries
            
        Returns:
            list: List of validated placeholders with page counts
        """
        validated = []
        input_dir = os.path.dirname(self.input_docx_path)
        
        print(f"\n🔍 VALIDATING {len(placeholders)} PDF references...")
        
        for i, placeholder in enumerate(placeholders):
            pdf_path_raw = placeholder['pdf_path_raw']
            
            print(f"\n   📋 Placeholder #{i+1} ({placeholder['type']}):")
            print(f"      • Raw path: {pdf_path_raw}")
            
            # Resolve relative path
            if not os.path.isabs(pdf_path_raw):
                pdf_path = os.path.join(input_dir, pdf_path_raw)
            else:
                pdf_path = pdf_path_raw
            
            pdf_path = os.path.abspath(pdf_path)
            print(f"      • Resolved path: {pdf_path}")
            
            # Check if file exists
            if not os.path.exists(pdf_path):
                print(f"      ❌ ERROR: PDF file not found")
                continue
              # Get page count
            try:
                pdf_doc = fitz.open(pdf_path)
                total_pages = pdf_doc.page_count
                pdf_doc.close()
                
                # Parse page specification if provided
                page_spec = placeholder.get('page_spec')
                if page_spec:
                    page_selection = self._parse_page_specification(page_spec)
                    if page_selection:
                        # Calculate which pages will actually be processed
                        pages = page_selection['pages']
                        open_range_start = page_selection.get('open_range_start')
                        
                        # Handle open-ended ranges
                        if open_range_start is not None:
                            pages.extend(range(open_range_start, total_pages))
                        
                        # Filter out pages that don't exist
                        valid_pages = [p for p in pages if 0 <= p < total_pages]
                        
                        if not valid_pages:
                            print(f"      ❌ ERROR: No valid pages in specification '{page_spec}' for {total_pages}-page PDF")
                            continue
                        
                        page_count = len(valid_pages)
                        print(f"      ✅ Valid PDF with {total_pages} total pages, using {page_count} specified pages: {[p+1 for p in valid_pages]}")
                        
                        # Store the actual pages to process
                        placeholder['selected_pages'] = valid_pages
                        placeholder['total_pages'] = total_pages
                    else:
                        # Invalid page specification, use all pages
                        page_count = total_pages
                        print(f"      ✅ Valid PDF with {page_count} page(s) (using all pages)")
                        placeholder['selected_pages'] = None
                        placeholder['total_pages'] = total_pages
                else:
                    # No page specification, use all pages
                    page_count = total_pages
                    print(f"      ✅ Valid PDF with {page_count} page(s)")
                    placeholder['selected_pages'] = None
                    placeholder['total_pages'] = total_pages
                  # Add resolved info to placeholder
                placeholder['pdf_path'] = pdf_path
                placeholder['page_count'] = page_count
                placeholder['index'] = len(validated)
                
                validated.append(placeholder)
                
            except Exception as e:
                print(f"      ❌ ERROR: Cannot read PDF file: {e}")
                continue
        
        print(f"\n✅ Validated {len(validated)}/{len(placeholders)} placeholders")
        return validated
    
    def _create_modified_document(self, overlay_placeholders, merge_placeholders):
        """
        Create a modified version of the document with insertion markers.
        
        Args:
            overlay_placeholders (list): Table-based placeholders
            merge_placeholders (list): Paragraph-based placeholders
            
        Returns:
            Document: Modified document with markers
        """
        print("\n🔧 Creating modified document...")
        doc = Document(self.input_docx_path)
        
        # Process merge placeholders (paragraph-based)
        if merge_placeholders:
            print(f"\n📄 Processing {len(merge_placeholders)} merge placeholders...")
            self._process_merge_placeholders(doc, merge_placeholders)
          # Process overlay placeholders (table-based)
        if overlay_placeholders:
            print(f"\n📦 Processing {len(overlay_placeholders)} overlay placeholders...")
            self._process_overlay_placeholders(doc, overlay_placeholders)
        
        print("✅ Document modification complete")
        return doc
    
    def _process_merge_placeholders(self, doc, merge_placeholders):
        """
        Process merge (paragraph-based) placeholders by adding markers and page breaks.
        
        Args:
            doc (Document): The document to modify
            merge_placeholders (list): List of merge placeholders
        """
        # Sort by paragraph index in reverse order to avoid index shifting
        sorted_placeholders = sorted(merge_placeholders, key=lambda x: x['paragraph_index'], reverse=True)
        
        for placeholder in sorted_placeholders:
            para_idx = placeholder['paragraph_index']
            page_count = placeholder['page_count']
            marker_text = f"%%MERGE_START_{placeholder['index']}%%"
            
            print(f"   📄 Processing merge placeholder #{placeholder['index']+1}:")
            print(f"      • Paragraph {para_idx}, {page_count} pages")
            print(f"      • Marker: {marker_text}")
              # Get the paragraph and replace content with marker + page break
            paragraph = doc.paragraphs[para_idx]
            paragraph.clear()
            
            # Add visible marker first
            marker_run = paragraph.add_run(marker_text)
            marker_run.font.size = Pt(12)
            marker_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            
            # Add page break after the marker to start PDF insertion on next page
            page_break_run = paragraph.add_run()
            page_break_run.add_break(WD_BREAK.PAGE)
            
            # Add the marker to the placeholder for processing
            placeholder['marker'] = marker_text
            
            print(f"      ✅ Added marker and page break (no placeholder pages)")
    
    def _process_overlay_placeholders(self, doc, overlay_placeholders):
        """
        Process table-based overlay placeholders by replacing table content with markers.
        
        Args:
            doc (Document): The document to modify
            overlay_placeholders (list): List of table-based overlay placeholders
        """
        print("   � Table-based overlay placeholder processing:")
        
        if not overlay_placeholders:
            return        
        # All overlay placeholders should be table-based now
        self._process_table_placeholders(doc, overlay_placeholders)
    
    def _process_table_placeholders(self, doc, table_placeholders):
        """
        Process table-based overlay placeholders by replacing table content with a simple marker.
        
        Args:
            doc (Document): The document to modify
            table_placeholders (list): List of table placeholders
        """        # Sort by table index in reverse order to avoid index shifting issues
        sorted_placeholders = sorted(table_placeholders, key=lambda x: x['table_index'], reverse=True)
        
        for placeholder in sorted_placeholders:
            table_idx = placeholder['table_index']
            page_count = placeholder['page_count']
            
            print(f"      • Processing table placeholder #{placeholder['index']+1}:")
            print(f"         • Table {table_idx}, {page_count} pages")
            
            try:
                # Get the table and extract its dimensions
                table = doc.tables[table_idx]
                cell = table.rows[0].cells[0]
                
                # Use stored table dimensions if available, otherwise calculate
                stored_coords = self.table_coordinates.get(table_idx, {})
                stored_dims = stored_coords.get('dimensions', {})
                
                if stored_dims and 'column_width_inches' in stored_dims and 'row_height_inches' in stored_dims:
                    # Use the stored dimensions from Word XML analysis
                    table_width_in = stored_dims['column_width_inches']
                    table_height_in = stored_dims['row_height_inches']
                    table_width_pts = table_width_in * 72  # Convert inches to points
                    table_height_pts = table_height_in * 72
                    print(f"         • Using stored table dimensions: {table_width_in:.2f} x {table_height_in:.2f} inches")
                else:
                    # Fallback: Extract table dimensions from docx properties
                    table_width_twips = table.columns[0].width
                    table_height_twips = sum(row.height for row in table.rows if row.height)
                    
                    # Convert to points (1 inch = 72 points, 1440 twips = 1 inch, so 20 twips = 1 point)
                    table_width_pts = table_width_twips / 20 if table_width_twips else 400  # Default width
                    table_height_pts = table_height_twips / 20 if table_height_twips else 200  # Default height
                    
                    # Apply reasonable limits (tables shouldn't be huge)
                    table_width_pts = min(table_width_pts, 600)  # Max ~8.3 inches
                    table_height_pts = min(table_height_pts, 800)  # Max ~11 inches
                    
                    # Ensure minimum size
                    table_width_pts = max(table_width_pts, 100)   # Min ~1.4 inches
                    table_height_pts = max(table_height_pts, 50)  # Min ~0.7 inches
                    
                    table_width_in = table_width_pts / 72
                    table_height_in = table_height_pts / 72
                    print(f"         • Calculated table dimensions: {table_width_in:.2f} x {table_height_in:.2f} inches")
                
                print(f"         • Final table dimensions: {table_width_pts:.1f} x {table_height_pts:.1f} points = {table_width_in:.2f} x {table_height_in:.2f} inches")
                
                # Clear the cell and add a simple top-left aligned marker
                cell.text = ""
                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                
                # Set paragraph alignment to left (top-left justification)
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Create the overlay marker
                marker = f"%%OVERLAY_START_{placeholder['index']:02d}%%"
                run = paragraph.add_run(marker)
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for visibility
                run.font.size = Pt(10)
                  # Store table dimensions and marker in the placeholder for overlay processing
                placeholder['marker'] = marker
                placeholder['table_width_pts'] = table_width_pts
                placeholder['table_height_pts'] = table_height_pts
                  # Handle multi-page PDFs by replicating cells instead of tables
                if page_count > 1:
                    print(f"         📋 Multi-page PDF detected ({page_count} pages), replicating cells...")
                    additional_markers = self._replicate_cells_for_multipage(
                        doc, table, placeholder, page_count - 1
                    )
                    placeholder['additional_markers'] = additional_markers
                    print(f"         ✅ Created {len(additional_markers)} additional cells")
                
                print(f"         ✅ Table {table_idx} updated with overlay marker and dimensions")
                
            except Exception as e:
                print(f"         ❌ Error modifying table {table_idx}: {e}")
                print(f"         📝 Falling back to simple marker...")
                # Fallback: add simple marker
                marker_text = f"%%OVERLAY_START_{placeholder['index']}%%"
                cell.text = marker_text
                placeholder['marker'] = marker_text
    def _replicate_cells_for_multipage(self, doc, table, placeholder, additional_pages):
        """
        Replicate table cells for multi-page PDF overlays by adding rows to the existing table.
        
        Args:
            doc (Document): The document to modify
            table: The table to add cells to
            placeholder (dict): The placeholder information
            additional_pages (int): Number of additional pages to create cells for
            
        Returns:
            list: List of markers for the additional cells
        """
        additional_markers = []
        
        try:
            # Add additional rows to the existing table
            for page_num in range(additional_pages):
                # Add a new row to the table
                new_row = table.add_row()
                cell = new_row.cells[0]
                
                # Set the row height to match the original row
                try:
                    if table.rows[0].height:
                        new_row.height = table.rows[0].height
                except:
                    pass  # Height setting might fail, continue without it
                
                # Configure the cell
                cell.text = ""
                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                
                # Set paragraph alignment to left (top-left justification)
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Create unique marker for this additional cell
                marker = f"%%OVERLAY_START_{placeholder['index']:02d}_PAGE_{page_num + 2:02d}%%"
                run = paragraph.add_run(marker)
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for visibility
                run.font.size = Pt(10)
                
                additional_markers.append({
                    'marker': marker,
                    'page_number': page_num + 2,  # Page 2, 3, 4, etc.
                    'table_width_pts': placeholder['table_width_pts'],
                    'table_height_pts': placeholder['table_height_pts']
                })
                
                print(f"           ✅ Created table row {page_num + 2} with marker: {marker}")
                
        except Exception as e:
            print(f"           ❌ Error during cell replication: {e}")
        
        return additional_markers

    def _convert_docx_to_pdf(self, input_path, output_path):
        """
        Convert a DOCX file to PDF using Microsoft Word automation.
        Uses tested conversion function with simplified parameters.
        
        Args:
            input_path (str): Path to the input DOCX file
            output_path (str): Path for the output PDF file
        """
        print(f"    Converting DOCX to PDF...")
        print(f"    Input: {input_path}")
        print(f"    Output: {output_path}")
        
        word = None
        doc = None
        
        try:
            # Get or create Word application
            try:
                word = win32.GetActiveObject("Word.Application")
                print("    ✓ Connected to existing Word instance")
            except:
                word = win32.Dispatch("Word.Application")
                word.Visible = False
                print("    ✓ Created new Word instance")
            
            # Open the document
            input_path = os.path.abspath(input_path)
            print(f"    Opening document: {input_path}")
            doc = word.Documents.Open(input_path)
            
            # Export parameters
            wdExportFormatPDF = 17
            wdExportOptimizeForPrint = 0
            wdExportCreateHeadingBookmarks = 1  # Create bookmarks from headings
            wdExportItem = 7  # Export entire document including markups
            
            output_path = os.path.abspath(output_path)
            print(f"    Exporting to PDF: {output_path}")
            
            # Export to PDF with minimal, tested parameters
            doc.ExportAsFixedFormat(
                OutputFileName=output_path,
                ExportFormat=wdExportFormatPDF,                OpenAfterExport=False,
                OptimizeFor=wdExportOptimizeForPrint,
                Item=wdExportItem,
                CreateBookmarks=wdExportCreateHeadingBookmarks
            )
            
            print(f"    ✓ Successfully converted '{os.path.basename(input_path)}' to PDF")
            
        except Exception as e:
            print(f"    ⚠ ERROR during Word conversion: {e}")
            raise
            
        finally:
            # Clean up Word objects using tested approach
            try:
                if 'doc' in locals() and doc:
                    doc.Close(SaveChanges=False)
                    print("    ✓ Document closed")
            except:
                pass
            
            try:
                if word and word.Documents.Count == 0:
                    word.Quit()
                    print("    ✓ Word application closed")
            except:
                pass
    
    def _overlay_pdfs(self, placeholders, source_pdf_path=None):
        """
        Overlay appendix PDFs onto the base PDF using simple marker positioning with table dimensions.
        
        Args:
            placeholders (list): List of placeholder dictionaries
            source_pdf_path (str): Path to source PDF (defaults to temp_pdf_path)
        """
        if source_pdf_path is None:
            source_pdf_path = self.temp_pdf_path
            
        print(f"    Opening base PDF: {source_pdf_path}")
        base_pdf = fitz.open(source_pdf_path)
        
        try:
            for placeholder in placeholders:
                pdf_path = placeholder['pdf_path']
                page_count = placeholder['page_count']
                index = placeholder['index']
                
                print(f"    Processing appendix {index + 1}: {os.path.basename(pdf_path)}")
                
                # Use the simplified approach with table dimensions
                overlay_rect, start_page_index = self._find_marker_and_calculate_rect_from_table(base_pdf, placeholder)
                
                if not overlay_rect:
                    print(f"      ⚠ WARNING: No markers found in PDF, skipping appendix")
                    continue
                
                print(f"      ✓ Found positioning markers on page {start_page_index + 1}")
                print(f"      📋 Overlay rectangle: ({overlay_rect.x0:.1f}, {overlay_rect.y0:.1f}) to ({overlay_rect.x1:.1f}, {overlay_rect.y1:.1f})")
                print(f"      📏 Overlay size: {overlay_rect.width:.1f} x {overlay_rect.height:.1f} points")                # Open the appendix PDF
                print(f"      Opening appendix PDF: {pdf_path}")
                appendix_pdf = fitz.open(pdf_path)
                
                try:
                    # Bake annotations into the PDF content to preserve them during overlay
                    print(f"      🔥 Baking annotations into PDF content...")
                    
                    # Check for annotations before baking
                    total_annotations = 0
                    for page_num in range(appendix_pdf.page_count):
                        page = appendix_pdf[page_num]
                        annots = list(page.annots())
                        total_annotations += len(annots)
                    
                    if total_annotations > 0:
                        print(f"        📝 Found {total_annotations} annotation(s), baking into content...")
                        appendix_pdf.bake(annots=True)
                        print(f"        ✓ Annotations baked into PDF content")
                    else:
                        print(f"        📝 No annotations found in PDF")                    # Determine which pages to overlay
                    selected_pages = placeholder.get('selected_pages')
                    if selected_pages is not None:
                        # Use specified pages
                        pages_to_overlay = selected_pages
                        print(f"        📄 Using selected pages: {[p+1 for p in pages_to_overlay]}")
                    else:
                        # Use all pages
                        pages_to_overlay = list(range(page_count))
                        print(f"        📄 Using all {page_count} pages")
                    
                    # Overlay each selected page of the appendix
                    for i, source_page_index in enumerate(pages_to_overlay):
                        if i == 0:
                            # First page: use the original marker and overlay rectangle
                            print(f"        Overlaying source page {source_page_index + 1} -> base page {start_page_index + 1} (position {i + 1}/{len(pages_to_overlay)})")
                            target_page = base_pdf[start_page_index]
                            print(f"        📌 Precise overlay within detected cell boundaries")
                            target_page.show_pdf_page(overlay_rect, appendix_pdf, source_page_index)
                        else:
                            # Additional pages: find the replicated table markers
                            additional_markers = placeholder.get('additional_markers', [])
                            if i - 1 < len(additional_markers):
                                marker_info = additional_markers[i - 1]
                                marker = marker_info['marker']
                                
                                print(f"        Overlaying source page {source_page_index + 1} -> searching for marker {marker} (position {i + 1}/{len(pages_to_overlay)})")
                                
                                # Find the marker for this additional page
                                additional_overlay_rect, additional_page_index = self._find_marker_and_calculate_rect_from_table_with_marker(
                                    base_pdf, marker, marker_info['table_width_pts'], marker_info['table_height_pts']
                                )
                                
                                if additional_overlay_rect:
                                    target_page = base_pdf[additional_page_index]
                                    print(f"        📌 Precise overlay in replicated table on page {additional_page_index + 1}")
                                    target_page.show_pdf_page(additional_overlay_rect, appendix_pdf, source_page_index)
                                else:
                                    print(f"        ❌ Could not find marker for position {i + 1}, skipping")
                            else:
                                print(f"        ❌ No replicated table found for position {i + 1}, skipping")
                    
                    print(f"      ✓ Appendix {index + 1} overlay complete")
                    
                finally:
                    appendix_pdf.close()            # Save the final PDF
            print(f"    Saving final PDF: {self.final_pdf_path}")
              # If we're overwriting the source file, save to temp first then replace
            if source_pdf_path == self.final_pdf_path:
                temp_output = f"{self.final_pdf_path}.tmp"
                base_pdf.save(temp_output)
                base_pdf.close()
                
                # Replace the original file
                import shutil
                shutil.move(temp_output, self.final_pdf_path)
                print("    ✓ Final PDF saved successfully")
                return  # Exit early since we already closed the PDF
            else:
                base_pdf.save(self.final_pdf_path)
                print("    ✓ Final PDF saved successfully")
                
        finally:
            if 'base_pdf' in locals() and not base_pdf.is_closed:
                base_pdf.close()
    def _insert_pdfs(self, merge_placeholders):
        """
        Insert PDFs at paragraph-based merge placeholders after marker positions.
        
        This method handles paragraph-based INSERT statements by:
        1. Finding the marker positions in the base PDF
        2. Removing the markers from the PDF
        3. Inserting the appendix PDF pages after the marker positions
        
        Args:
            merge_placeholders (list): List of paragraph-based placeholder dictionaries
        """
        print(f"    Opening base PDF for merge insertions: {self.temp_pdf_path}")
        base_pdf = fitz.open(self.temp_pdf_path)
        
        try:
            # Process placeholders in reverse order to avoid page index shifting
            sorted_placeholders = sorted(merge_placeholders, key=lambda x: x.get('index', 0), reverse=True)
            
            for placeholder in sorted_placeholders:
                pdf_path = placeholder['pdf_path']
                marker = placeholder.get('marker', '')
                index = placeholder.get('index', 0)
                
                print(f"    Processing merge appendix {index + 1}: {os.path.basename(pdf_path)}")
                print(f"      🔍 Searching for marker: {marker}")
                
                # Find the marker position in the base PDF
                marker_info = self._find_marker_position(base_pdf, marker)
                
                if not marker_info:
                    print(f"      ❌ Marker not found in PDF, skipping appendix")
                    continue
                
                start_page_index = marker_info['page_index']
                
                print(f"      ✓ Found marker on page {start_page_index + 1}")
                
                # Remove the marker from the PDF page
                print(f"      🧹 Removing marker from page {start_page_index + 1}")
                marker_page = base_pdf[start_page_index]
                text_instances = marker_page.search_for(marker)
                if text_instances:
                    for inst in text_instances:
                        # Create a white rectangle to cover the marker text
                        marker_page.add_redact_annot(inst, fill=(1, 1, 1))  # White fill
                    marker_page.apply_redactions()
                    print(f"      ✓ Marker removed from page {start_page_index + 1}")
                
                print(f"      📄 Will insert {len(pages_to_insert) if 'pages_to_insert' in locals() else 'PDF'} page(s) after page {start_page_index + 1}")
                
                # Open the appendix PDF
                print(f"      Opening appendix PDF: {pdf_path}")
                appendix_pdf = fitz.open(pdf_path)
                
                try:
                    # Bake annotations into the PDF content to preserve them
                    print(f"      🔥 Baking annotations into PDF content...")
                    
                    # Check for annotations before baking
                    total_annotations = 0
                    for page_num in range(appendix_pdf.page_count):
                        page = appendix_pdf[page_num]
                        annots = list(page.annots())
                        total_annotations += len(annots)
                    
                    if total_annotations > 0:
                        print(f"        📝 Found {total_annotations} annotation(s), baking into content...")
                        appendix_pdf.bake(annots=True)
                        print(f"        ✓ Annotations baked into PDF content")
                    else:
                        print(f"        📝 No annotations found in PDF")
                    
                    # Determine which pages to insert
                    selected_pages = placeholder.get('selected_pages')
                    if selected_pages is not None:
                        # Use specified pages
                        pages_to_insert = selected_pages
                        print(f"        📄 Using selected pages: {[p+1 for p in pages_to_insert]}")
                    else:
                        # Use all pages
                        pages_to_insert = list(range(appendix_pdf.page_count))
                        print(f"        📄 Using all {appendix_pdf.page_count} pages")
                    
                    # Insert the appendix pages after the marker page (no page removal)
                    insert_start_position = start_page_index + 1  # Insert after the marker page
                    print(f"        📥 Inserting {len(pages_to_insert)} appendix page(s) starting at position {insert_start_position + 1}")
                    for i, source_page_index in enumerate(pages_to_insert):
                        insert_position = insert_start_position + i
                        print(f"          • Inserting source page {source_page_index + 1} at position {insert_position + 1}")
                        base_pdf.insert_pdf(appendix_pdf, from_page=source_page_index, to_page=source_page_index, start_at=insert_position)
                    
                    print(f"      ✓ Merge appendix {index + 1} insertion complete")
                    
                finally:
                    appendix_pdf.close()
            
            # Save the final PDF (or intermediate PDF if overlays will follow)
            print(f"    Saving PDF with merges: {self.final_pdf_path}")
            base_pdf.save(self.final_pdf_path)
            print("    ✓ PDF with merges saved successfully")
            
        finally:
            base_pdf.close()
    
    def _find_marker_and_calculate_rect_from_table(self, pdf_doc, placeholder):
        """
        Find the overlay marker and calculate the overlay rectangle using table dimensions.
        
        Args:
            pdf_doc: The PDF document to search
            placeholder (dict): Placeholder information with marker and table dimensions
            
        Returns:
            tuple: (overlay_rect, page_index) or (None, None) if not found
        """
        marker = placeholder.get('marker')
        if not marker:
            print(f"      ❌ No marker found in placeholder")
            return None, None
        
        print(f"      🔍 Searching for marker: {marker}")
        
        # Search for the marker in the PDF
        marker_info = self._find_marker_position(pdf_doc, marker)
        
        if not marker_info:
            print(f"      ❌ Marker not found in PDF")
            return None, None
        
        start_page_index = marker_info['page_index']
        marker_rect = marker_info['rect']
          # Convert points to inches for easier reading (72 points = 1 inch)
        marker_x_in = marker_rect.x0 / 72
        marker_y_in = marker_rect.y0 / 72
        marker_width_in = marker_rect.width / 72
        marker_height_in = marker_rect.height / 72
        
        print(f"      📍 Marker found at: ({marker_rect.x0:.1f}, {marker_rect.y0:.1f}) points = ({marker_x_in:.2f}, {marker_y_in:.2f}) inches")
        print(f"      📍 Marker size: {marker_rect.width:.1f} x {marker_rect.height:.1f} points = {marker_width_in:.2f} x {marker_height_in:.2f} inches")
        
        # Get table dimensions from placeholder
        table_width_pts = placeholder.get('table_width_pts', 400)
        table_height_pts = placeholder.get('table_height_pts', 200)
        table_width_in = table_width_pts / 72
        table_height_in = table_height_pts / 72
        
        print(f"      📏 Table dimensions: {table_width_pts:.1f} x {table_height_pts:.1f} points = {table_width_in:.2f} x {table_height_in:.2f} inches")
          # Calculate overlay rectangle using marker position as top-left corner
        # and adding table dimensions for bottom-right corner
        overlay_rect = fitz.Rect(
            marker_rect.x0,                           # left (marker x position)
            marker_rect.y0,                           # top (marker y position)  
            marker_rect.x0 + table_width_pts,         # right (left + table width)
            marker_rect.y0 + table_height_pts         # bottom (top + table height)
        )
        
        # Convert overlay rectangle to inches for easier reading
        overlay_x_in = overlay_rect.x0 / 72
        overlay_y_in = overlay_rect.y0 / 72
        overlay_width_in = overlay_rect.width / 72
        overlay_height_in = overlay_rect.height / 72
        
        print(f"      📐 Calculated overlay rectangle:")
        print(f"         • Points: ({overlay_rect.x0:.1f}, {overlay_rect.y0:.1f}) to ({overlay_rect.x1:.1f}, {overlay_rect.y1:.1f})")
        print(f"         • Inches: ({overlay_x_in:.2f}, {overlay_y_in:.2f}) to ({overlay_x_in + overlay_width_in:.2f}, {overlay_y_in + overlay_height_in:.2f})")
        print(f"         • Size: {overlay_rect.width:.1f} x {overlay_rect.height:.1f} points = {overlay_width_in:.2f} x {overlay_height_in:.2f} inches")
        
        # Remove the marker text from the page
        page = pdf_doc[start_page_index]
        page.add_redact_annot(marker_rect, fill=(1, 1, 1))  # White fill
        page.apply_redactions()
        print(f"      ✓ Removed marker text from page {start_page_index + 1}")
        
        return overlay_rect, start_page_index    
    def _find_marker_position(self, pdf_doc, marker):
        """
        Find the position of a marker in the PDF document.
        
        Args:
            pdf_doc: The PDF document to search
            marker (str): The marker text to find
            
        Returns:
            dict: Dictionary with page_index and rect, or None if not found
        """
        for page_num in range(pdf_doc.page_count):
            page = pdf_doc[page_num]
            text_instances = page.search_for(marker)
            
            if text_instances:                # Return the first instance found
                return {
                    'page_index': page_num,
                    'rect': text_instances[0]
                }
        
        return None
    
    def _find_marker_and_calculate_rect_from_table_with_marker(self, pdf_doc, marker, table_width_pts, table_height_pts):
        """
        Find a specific marker and calculate overlay rectangle using provided table dimensions.
        
        Args:
            pdf_doc: The PDF document to search
            marker (str): The marker text to find
            table_width_pts (float): Table width in points
            table_height_pts (float): Table height in points
            
        Returns:
            tuple: (overlay_rect, page_index) or (None, None) if not found
        """
        print(f"      🔍 Searching for marker: {marker}")
        
        # Search for the marker in the PDF
        marker_info = self._find_marker_position(pdf_doc, marker)
        
        if not marker_info:
            print(f"      ❌ Marker not found in PDF")
            return None, None
        
        start_page_index = marker_info['page_index']
        marker_rect = marker_info['rect']
        
        # Convert points to inches for easier reading
        marker_x_in = marker_rect.x0 / 72
        marker_y_in = marker_rect.y0 / 72
        marker_width_in = marker_rect.width / 72
        marker_height_in = marker_rect.height / 72
        
        print(f"      📍 Marker found at: ({marker_rect.x0:.1f}, {marker_rect.y0:.1f}) points = ({marker_x_in:.2f}, {marker_y_in:.2f}) inches")
        print(f"      📍 Marker size: {marker_rect.width:.1f} x {marker_rect.height:.1f} points = {marker_width_in:.2f} x {marker_height_in:.2f} inches")
        
        table_width_in = table_width_pts / 72
        table_height_in = table_height_pts / 72
        
        print(f"      📏 Table dimensions: {table_width_pts:.1f} x {table_height_pts:.1f} points = {table_width_in:.2f} x {table_height_in:.2f} inches")
        
        # Calculate overlay rectangle using marker position as top-left corner
        overlay_rect = fitz.Rect(
            marker_rect.x0,                           # left (marker x position)
            marker_rect.y0,                           # top (marker y position)  
            marker_rect.x0 + table_width_pts,         # right (left + table width)
            marker_rect.y0 + table_height_pts         # bottom (top + table height)
        )
        
        # Convert overlay rectangle to inches for easier reading
        overlay_x_in = overlay_rect.x0 / 72
        overlay_y_in = overlay_rect.y0 / 72
        overlay_width_in = overlay_rect.width / 72
        overlay_height_in = overlay_rect.height / 72
        
        print(f"      📐 Calculated overlay rectangle:")
        print(f"         • Points: ({overlay_rect.x0:.1f}, {overlay_rect.y0:.1f}) to ({overlay_rect.x1:.1f}, {overlay_rect.y1:.1f})")
        print(f"         • Inches: ({overlay_x_in:.2f}, {overlay_y_in:.2f}) to ({overlay_x_in + overlay_width_in:.2f}, {overlay_y_in + overlay_height_in:.2f})")
        print(f"         • Size: {overlay_rect.width:.1f} x {overlay_rect.height:.1f} points = {overlay_width_in:.2f} x {overlay_height_in:.2f} inches")
        
        # Remove the marker text from the page
        page = pdf_doc[start_page_index]
        page.add_redact_annot(marker_rect, fill=(1, 1, 1))  # White fill
        page.apply_redactions()
        print(f"      ✓ Removed marker text from page {start_page_index + 1}")
        
        return overlay_rect, start_page_index

    def _cleanup(self):
        """
        Clean up temporary files created during the process.
        Skips cleanup if keep_temp flag is set.
        """
        if self.keep_temp:
            print("\nKeeping temporary files for debugging:")
            files_to_keep = [self.temp_docx_path, self.temp_pdf_path]
            for file_path in files_to_keep:
                if os.path.exists(file_path):
                    print(f"  ✓ Kept: {file_path}")
                else:
                    print(f"  - Not found: {os.path.basename(file_path)}")
            return
        
        print("\nCleaning up temporary files...")
        
        files_to_remove = [self.temp_docx_path, self.temp_pdf_path]
        
        for file_path in files_to_remove:
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                    print(f"  ✓ Removed: {os.path.basename(file_path)}")
                except Exception as e:
                    print(f"  ⚠ Could not remove {os.path.basename(file_path)}: {e}")
            else:
                print(f"  - Not found: {os.path.basename(file_path)}")


def main():
    """
    Main execution function with command-line argument parsing.
    """
    parser = argparse.ArgumentParser(
        description="Python PDF Report Compiler - Combine Word documents with PDF appendices",
        formatter_class=argparse.RawDescriptionHelpFormatter,        epilog="""
Examples:
  python report_compiler.py report.docx final_report.pdf
  python report_compiler.py "C:\\Reports\\my_report.docx" "C:\\Output\\final.pdf"

Placeholder formats in Word document:
  Table-based overlays (in single-cell tables):
    [[OVERLAY: appendices/sketch.pdf, page=1]]
    [[OVERLAY: C:\\Shared\\calculation.pdf]]
  
  Paragraph-based merges (in standalone paragraphs):
    [[INSERT: appendices/full_report.pdf]]
    [[INSERT: C:\\Shared\\analysis.pdf:1-5]]        """
    )
    
    parser.add_argument(
        'input_file',
        help='Path to the input Word document (.docx)'
    )
    
    parser.add_argument(
        'output_file', 
        help='Path for the output PDF file (.pdf)'
    )
    
    parser.add_argument(
        '--keep-temp',
        action='store_true',
        help='Keep temporary files for debugging purposes'
    )
    
    args = parser.parse_args()
    
    # Validate input file
    if not args.input_file.lower().endswith('.docx'):
        print("ERROR: Input file must be a .docx file")
        return 1
    
    if not os.path.exists(args.input_file):
        print(f"ERROR: Input file not found: {args.input_file}")
        return 1
    
    # Validate output file
    if not args.output_file.lower().endswith('.pdf'):
        print("ERROR: Output file must be a .pdf file")
        return 1
      # Create output directory if it doesn't exist
    output_dir = os.path.dirname(os.path.abspath(args.output_file))
    if output_dir and not os.path.exists(output_dir):
        try:
            os.makedirs(output_dir)
            print(f"Created output directory: {output_dir}")
        except Exception as e:
            print(f"ERROR: Could not create output directory: {e}")
            return 1
    
    try:
        # Create and run the compiler
        compiler = ReportCompiler(args.input_file, args.output_file, keep_temp=args.keep_temp)
        compiler.run()
        return 0
        
    except Exception as e:
        print(f"\nERROR: Compilation failed: {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == "__main__":
    exit(main())
