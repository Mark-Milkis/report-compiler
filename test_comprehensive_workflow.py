#!/usr/bin/env python3
"""
End-to-end test of IMAGE functionality with the complete workflow.
"""

import os
import sys
import tempfile

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from report_compiler.document.placeholder_parser import PlaceholderParser
from report_compiler.document.docx_processor import DocxProcessor
from report_compiler.utils.validators import Validators

def create_comprehensive_test():
    """Create a comprehensive test with IMAGE and OVERLAY placeholders."""
    from PIL import Image
    from docx import Document
    from docx.shared import Inches
    
    # Create test images
    img1 = Image.new('RGB', (400, 300), color='red')
    img1_path = '/tmp/red_image.png'
    img1.save(img1_path)
    
    img2 = Image.new('RGB', (200, 150), color='blue') 
    img2_path = '/tmp/blue_image.jpg'
    img2.save(img2_path)
    
    # Create test document
    doc = Document()
    doc.add_heading('Mixed Placeholder Test Document', 0)
    
    # Add some content
    doc.add_paragraph('This document tests both IMAGE and OVERLAY placeholders.')
    
    # Test 1: IMAGE placeholder
    doc.add_heading('Test 1: IMAGE Placeholder', level=1)
    doc.add_paragraph('The following table contains an image:')
    
    table1 = doc.add_table(rows=1, cols=1)
    table1.style = 'Table Grid'
    cell1 = table1.cell(0, 0)
    cell1.text = f"[[IMAGE: {img1_path}]]"
    table1.columns[0].width = Inches(4)
    table1.rows[0].height = Inches(3)
    
    # Test 2: IMAGE with parameters
    doc.add_heading('Test 2: IMAGE with Size Parameters', level=1)
    table2 = doc.add_table(rows=1, cols=1)
    table2.style = 'Table Grid'
    cell2 = table2.cell(0, 0)
    cell2.text = f"[[IMAGE: {img2_path}, width=2in, height=1.5in]]"
    table2.columns[0].width = Inches(3)
    table2.rows[0].height = Inches(2)
    
    # Test 3: OVERLAY placeholder (should still work)
    doc.add_heading('Test 3: OVERLAY Placeholder (Legacy)', level=1)
    doc.add_paragraph('This will show as a marker since no PDF processing happens in this test:')
    table3 = doc.add_table(rows=1, cols=1)
    table3.style = 'Table Grid'
    cell3 = table3.cell(0, 0)
    cell3.text = "[[OVERLAY: /tmp/sample.pdf]]"
    table3.columns[0].width = Inches(3)
    table3.rows[0].height = Inches(2)
    
    test_docx_path = '/tmp/comprehensive_test.docx'
    doc.save(test_docx_path)
    
    return test_docx_path, [img1_path, img2_path]

def test_complete_workflow():
    """Test the complete workflow with IMAGE placeholders."""
    print("Running comprehensive IMAGE workflow test...\n")
    
    # Create test files
    test_docx_path, image_paths = create_comprehensive_test()
    output_docx_path = '/tmp/comprehensive_output.docx'
    
    print(f"Created test document: {test_docx_path}")
    print(f"Created test images: {image_paths}")
    
    # Step 1: Parse placeholders
    print("\n1. Parsing placeholders...")
    parser = PlaceholderParser()
    placeholders = parser.find_all_placeholders(test_docx_path)
    
    print(f"   Found {placeholders['total']} total placeholders")
    print(f"   Table placeholders: {len(placeholders['table'])}")
    print(f"   Paragraph placeholders: {len(placeholders['paragraph'])}")
    
    for i, placeholder in enumerate(placeholders['table']):
        subtype = placeholder.get('subtype', 'overlay')
        print(f"   Table {i}: {subtype} - {placeholder['file_path']}")
    
    # Step 2: Validate placeholders
    print("\n2. Validating placeholders...")
    validators = Validators()
    validation_result = validators.validate_placeholders(
        placeholders['table'] + placeholders['paragraph'], 
        '/tmp'
    )
    
    print(f"   Validation valid: {validation_result['valid']}")
    if validation_result['errors']:
        print(f"   Errors: {validation_result['errors']}")
    if validation_result['warnings']:
        print(f"   Warnings: {validation_result['warnings']}")
    
    # Step 3: Process the document
    print("\n3. Processing document...")
    processor = DocxProcessor()
    metadata = processor.create_modified_docx(test_docx_path, placeholders, output_docx_path)
    
    if metadata is not None:
        print(f"   ✓ Document processed successfully")
        print(f"   ✓ Output saved to: {output_docx_path}")
        
        # Check output file
        if os.path.exists(output_docx_path):
            size = os.path.getsize(output_docx_path)
            print(f"   ✓ Output file size: {size} bytes")
            
            # Verify content
            try:
                from docx import Document
                output_doc = Document(output_docx_path)
                
                image_count = 0
                marker_count = 0
                
                for i, table in enumerate(output_doc.tables):
                    if len(table.rows) == 1 and len(table.columns) == 1:
                        cell = table.cell(0, 0)
                        
                        # Check for images
                        inline_shapes = cell.paragraphs[0]._element.xpath('.//a:blip')
                        if inline_shapes:
                            image_count += 1
                            print(f"   ✓ Table {i} contains an image")
                        elif '%%OVERLAY_START_' in cell.text:
                            marker_count += 1
                            print(f"   ✓ Table {i} contains overlay marker: {cell.text}")
                        else:
                            print(f"   ? Table {i} content: '{cell.text}'")
                
                print(f"   ✓ Found {image_count} images and {marker_count} overlay markers")
                
            except Exception as e:
                print(f"   ✗ Error verifying output: {e}")
        else:
            print(f"   ✗ Output file was not created")
    else:
        print(f"   ✗ Document processing failed")

if __name__ == "__main__":
    test_complete_workflow()
    print("\n✓ Comprehensive workflow test completed!")