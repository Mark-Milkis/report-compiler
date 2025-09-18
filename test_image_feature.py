#!/usr/bin/env python3
"""
Simple test script to verify IMAGE tag functionality.
"""

import os
import sys
import tempfile
from pathlib import Path

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from report_compiler.document.placeholder_parser import PlaceholderParser
from report_compiler.core.config import Config

def test_image_regex():
    """Test that the IMAGE regex pattern works correctly."""
    print("Testing IMAGE regex pattern...")
    
    test_cases = [
        "[[IMAGE: test.png]]",
        "[[IMAGE: path/to/image.jpg]]",
        "[[IMAGE: image.png, width=2in]]",
        "[[IMAGE: image.jpg, width=2in, height=1.5in]]",
        "[[image: test.PNG]]",  # Case insensitive
    ]
    
    for test_case in test_cases:
        match = Config.IMAGE_REGEX.search(test_case)
        if match:
            path = match.group(1).strip()
            params = match.group(2)
            print(f"  ✓ '{test_case}' -> path='{path}', params='{params}'")
        else:
            print(f"  ✗ '{test_case}' did not match")

def test_image_parameter_parsing():
    """Test that image parameters are parsed correctly."""
    print("\nTesting IMAGE parameter parsing...")
    
    parser = PlaceholderParser()
    
    test_cases = [
        (None, {}),
        ("width=2in", {"width": "2in"}),
        ("height=1.5in", {"height": "1.5in"}),
        ("width=2in, height=1.5in", {"width": "2in", "height": "1.5in"}),
        ("width=100px, height=200px", {"width": "100px", "height": "200px"}),
    ]
    
    for params_string, expected in test_cases:
        result = parser._parse_image_parameters(params_string)
        if result == expected:
            print(f"  ✓ '{params_string}' -> {result}")
        else:
            print(f"  ✗ '{params_string}' -> {result}, expected {expected}")

def create_test_image():
    """Create a simple test image for testing."""
    try:
        from PIL import Image
        
        # Create a simple 100x100 red square
        img = Image.new('RGB', (100, 100), color='red')
        test_image_path = '/tmp/test_image.png'
        img.save(test_image_path)
        print(f"\nCreated test image: {test_image_path}")
        return test_image_path
    except ImportError:
        print("\nPIL not available, skipping image creation")
        return None

def create_test_docx():
    """Create a simple test DOCX with IMAGE placeholders."""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading('Test Document with IMAGE Tags', 0)
        
        # Add a paragraph
        doc.add_paragraph('This document contains IMAGE placeholders for testing.')
        
        # Create a simple single-cell table with IMAGE placeholder
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        cell = table.cell(0, 0)
        cell.text = "[[IMAGE: /tmp/test_image.png]]"
        
        # Set table dimensions
        table.columns[0].width = Inches(3)
        table.rows[0].height = Inches(2)
        
        test_docx_path = '/tmp/test_document.docx'
        doc.save(test_docx_path)
        print(f"Created test document: {test_docx_path}")
        return test_docx_path
    except ImportError:
        print("python-docx not available, skipping document creation")
        return None

def test_placeholder_detection():
    """Test that IMAGE placeholders are detected in a DOCX document."""
    print("\nTesting IMAGE placeholder detection...")
    
    test_image_path = create_test_image()
    test_docx_path = create_test_docx()
    
    if not test_docx_path:
        print("  Skipping test - could not create test document")
        return
    
    parser = PlaceholderParser()
    placeholders = parser.find_all_placeholders(test_docx_path)
    
    print(f"  Found {placeholders['total']} total placeholders")
    print(f"  Table placeholders: {len(placeholders['table'])}")
    print(f"  Paragraph placeholders: {len(placeholders['paragraph'])}")
    
    for i, placeholder in enumerate(placeholders['table']):
        print(f"    Table {i}: {placeholder}")

if __name__ == "__main__":
    print("Testing IMAGE tag feature implementation...\n")
    
    test_image_regex()
    test_image_parameter_parsing()
    test_placeholder_detection()
    
    print("\n✓ All tests completed!")